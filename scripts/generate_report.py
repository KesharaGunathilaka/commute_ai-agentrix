"""
Generates the CommuteAI Project Report as a DOCX file.
Run: python scripts/generate_report.py
Output: CommuteAI_Project_Report.docx
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "CommuteAI_Project_Report.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
C_HEADER     = RGBColor(0x1A, 0x37, 0x5E)   # Deep navy
C_ACCENT     = RGBColor(0x21, 0x7D, 0xBB)   # Teal blue
C_LIGHT_BG   = RGBColor(0xEA, 0xF4, 0xFB)   # Light sky
C_DARK_TEXT  = RGBColor(0x1C, 0x1C, 0x1C)   # Near-black
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_TABLE_HDR  = RGBColor(0x1A, 0x37, 0x5E)   # Same as header
C_TABLE_ALT  = RGBColor(0xF0, 0xF7, 0xFC)   # Very light blue alt row
C_BORDER     = RGBColor(0xB0, 0xCC, 0xDC)
C_WARN       = RGBColor(0xC0, 0x39, 0x2B)
C_GREEN      = RGBColor(0x1A, 0x7A, 0x4A)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    hex_color = str(rgb)  # e.g. "1A375E"
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14) if level == 1 else Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = C_HEADER
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = C_ACCENT
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = C_DARK_TEXT
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = C_DARK_TEXT
    return p


def add_body(doc, text, size=10, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = C_DARK_TEXT
    return p


def add_bullet(doc, text, size=9.5, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = C_DARK_TEXT
    return p


def add_code_block(doc, text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 90)
    run.font.size = Pt(7)
    run.font.color.rgb = C_BORDER


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, C_TABLE_HDR)
        _cell_text(cell, h, bold=True, size=9, color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row in enumerate(rows):
        bg = C_TABLE_ALT if ri % 2 == 0 else C_WHITE
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            _set_cell_bg(cell, bg)
            _cell_text(cell, str(val), size=9)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ── Document ───────────────────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CommuteAI")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = C_HEADER

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Disruption-Aware Commute Agent for Sri Lanka Public Transport")
    r2.font.size = Pt(14)
    r2.font.color.rgb = C_ACCENT

    doc.add_paragraph()
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Project Report")
    r3.bold = True
    r3.font.size = Pt(18)
    r3.font.color.rgb = C_DARK_TEXT

    doc.add_paragraph()
    doc.add_paragraph()

    meta_lines = [
        ("Hackathon",    "AGENTRIX 2026"),
        ("Team",         "Team 23 — AlphaZero"),
        ("Repository",   "AGENTRIX26-TEAM23-AlphaZero"),
        ("Version",      "0.1.0"),
        ("Python",       "3.11+"),
        ("Framework",    "LangGraph · LangChain · Streamlit"),
    ]
    tbl = doc.add_table(rows=len(meta_lines), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta_lines):
        row = tbl.rows[i]
        _set_cell_bg(row.cells[0], C_TABLE_HDR)
        _cell_text(row.cells[0], k, bold=True, size=10, color=C_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_bg(row.cells[1], C_TABLE_ALT)
        _cell_text(row.cells[1], v, size=10)
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1.  Project Overview",
        "2.  System Architecture",
        "3.  Architecture Diagram",
        "4.  Technology Stack",
        "5.  Domain Models & ER Diagram",
        "6.  Agent Graph — Node Reference",
        "7.  Use Case Diagram",
        "8.  API Reference",
        "9.  Data Schema Reference",
        "10. Configuration & Prompts",
        "11. NLU Pipeline",
        "12. Ride-Hailing Integration",
        "13. UI & User Experience",
        "14. Execution Flow Walkthrough",
        "15. Design Decisions & Trade-offs",
        "16. Deployment Guide",
    ]
    for line in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(10)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Project Overview", 1)
    add_body(doc,
        "CommuteAI is a disruption-aware, multilingual commute planning agent built for Sri Lanka "
        "public transport. It integrates live route discovery (Google Maps Directions API), real-time "
        "train schedules (scraped from trainschedule.lk), bus timetable lookup, intelligent disruption "
        "detection, automatic replanning, ride-hailing fallback, and bilingual text-to-speech output. "
        "The system is orchestrated as a LangGraph state machine and served through a conversational "
        "Streamlit chat interface.")

    add_heading(doc, "Key Capabilities", 2)
    capabilities = [
        "Multilingual intent parsing — Sinhala, Tamil, English via Groq LLM (llama-3.3-70b-versatile)",
        "Live route discovery — Google Maps Directions API (transit mode, bus + train alternatives)",
        "Bus timetable enrichment — next scheduled departure from local bus_timetables.json",
        "Train schedule enrichment — live scrape of trainschedule.lk with station name normalization",
        "Multi-criteria route ranking — deadline compliance, departure timing, mode preference, transfers",
        "Disruption detection — segment-based matching against active entries in disruptions.json",
        "Bounded auto-replanning — up to 2 alternative routes before graceful fallback",
        "Ride-hailing fallback — RideService quotes (bike/tuk/car) when transit is unavailable or late",
        "Bilingual responses — native language + English translation always generated together",
        "Text-to-speech audio — per-message Read Aloud button using Google TTS (gTTS)",
        "Conversational intake — multi-turn collection of origin, destination, time, arrival deadline",
        "Agent execution trace — visible in the UI for full transparency and demo purposes",
    ]
    for c in capabilities:
        add_bullet(doc, c)

    add_heading(doc, "Problem Statement", 2)
    add_body(doc,
        "Sri Lanka commuters face frequent unannounced train delays and cancellations, with no "
        "integrated system that combines schedule data, disruption alerts, and alternative routing "
        "in their native languages. CommuteAI addresses this gap by acting as a personal commute "
        "assistant that proactively detects disruptions, replans journeys, and communicates clearly "
        "in Sinhala, Tamil, or English.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. System Architecture", 1)
    add_body(doc,
        "CommuteAI is built on a layered architecture with six distinct layers. Each layer has a "
        "single responsibility and communicates only through well-defined interfaces.")

    layers = [
        ("Presentation Layer",   "Streamlit chat UI (app.py) — conversational intake, route cards, ranked routes, Uber cards, Read Aloud button"),
        ("API Layer",            "FastAPI REST API — /api/v1/query, /api/v1/health, /api/v1/disruptions/{id}/activate"),
        ("Orchestration Layer",  "LangGraph StateGraph — 8 nodes, conditional edges, bounded replan loop"),
        ("NLU Layer",            "Groq LLM (parse_query, parse_update, respond_*) — intent parsing and bilingual response generation"),
        ("Data & Tools Layer",   "Google Maps API, trainschedule.lk scraper, bus_timetables.json, disruptions.json, RideService"),
        ("Persistence Layer",    "ChromaDB vector store (routes), JSON files (routes, disruptions, stations, bus_timetables), TTL cache"),
    ]
    make_table(doc, ["Layer", "Components"], layers, col_widths=[5, 12])

    add_heading(doc, "Component Interaction", 2)
    add_body(doc,
        "The Streamlit UI accumulates user intent over multiple conversational turns, then calls "
        "run_commute_agent_from_intent() which invokes the compiled LangGraph graph. The graph "
        "threads state through all 8 nodes sequentially (with conditional branching for disruptions). "
        "Each node reads specific fields from AgentState, performs its work (LLM call, API call, "
        "file lookup, or scrape), then writes results back into AgentState. The final state is "
        "returned to the UI for rendering.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. ARCHITECTURE DIAGRAM
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Architecture Diagram", 1)
    add_body(doc,
        "The following diagram illustrates the full system architecture showing all layers, "
        "components, external services, and the data flow between them.")

    arch_diagram = """\
┌─────────────────────────────────────────────────────────────────────────────┐
│                        COMMUTEAI — SYSTEM ARCHITECTURE                       │
└─────────────────────────────────────────────────────────────────────────────┘

  ┌──────────────────────────────────────────────────────────────────────────┐
  │  USER INTERFACES                                                          │
  │  ┌────────────────────────────────┐   ┌──────────────────────────────┐   │
  │  │  Streamlit Chat UI (app.py)    │   │  FastAPI REST API             │   │
  │  │  • Conversational intake       │   │  POST /api/v1/query           │   │
  │  │  • Route cards & ranked routes │   │  GET  /api/v1/health          │   │
  │  │  • Uber cards                  │   │  POST /disruptions/{id}       │   │
  │  │  • Read Aloud (gTTS button)    │   │  /activate                    │   │
  │  │  • Agent trace panel           │   │                               │   │
  │  └────────────────┬───────────────┘   └───────────────┬──────────────┘   │
  └───────────────────┼───────────────────────────────────┼──────────────────┘
                      │  run_commute_agent_from_intent()  │  run_commute_agent()
                      └─────────────────┬─────────────────┘
                                        ▼
  ┌──────────────────────────────────────────────────────────────────────────┐
  │  LANGGRAPH ORCHESTRATION LAYER (AgentState TypedDict)                     │
  │                                                                            │
  │  [START]                                                                   │
  │     │                                                                      │
  │  ┌──▼──────┐   ┌─────────┐   ┌──────────┐   ┌────────┐   ┌──────────┐   │
  │  │ PLANNER │──▶│ BUS_RAG │──▶│TRAIN_RAG │──▶│ RANKER │──▶│  UBER   │   │
  │  └─────────┘   └─────────┘   └──────────┘   └────────┘   └────┬─────┘   │
  │      │ (error/no route)                                         │          │
  │      └────────────────────────────────────┐                    ▼          │
  │                                           │              ┌─────────┐      │
  │                                           │              │ MONITOR │      │
  │                                           │              └────┬────┘      │
  │                                           │         clear │   │ disrupted │
  │                                           │               ▼   ▼           │
  │                                           │  ┌──────────┐ ┌───────────┐  │
  │                                           └─▶│ RESPONDER│ │REPLANNER  │  │
  │                                              └──────────┘ └──────┬────┘  │
  │                                                   │              │(loop)  │
  │                                                 [END]      back to MONITOR│
  └──────────────────────────────────────────────────────────────────────────┘
                                        │
              ┌─────────────────────────┼──────────────────────────┐
              ▼                         ▼                           ▼
  ┌─────────────────┐      ┌──────────────────────┐    ┌─────────────────────┐
  │  GROQ LLM API   │      │  GOOGLE MAPS API      │    │  trainschedule.lk   │
  │ llama-3.3-70b   │      │  Directions (transit) │    │  (web scraping)     │
  │ • parse_query   │      │  Bus + train routes   │    │  Train timetables   │
  │ • parse_update  │      │  Per-leg details      │    │  390 stations       │
  │ • respond_*     │      └──────────────────────┘    └─────────────────────┘
  │ • map_station   │
  └─────────────────┘
              │
              ▼
  ┌──────────────────────────────────────────────────────────────────────────┐
  │  DATA & PERSISTENCE LAYER                                                  │
  │  ┌────────────────┐  ┌────────────────┐  ┌────────────┐  ┌────────────┐  │
  │  │  ChromaDB      │  │ disruptions    │  │ bus        │  │ stations   │  │
  │  │  (vector store)│  │ .json (feed)   │  │_timetables │  │ .json      │  │
  │  │  routes embed  │  │ D001/D002      │  │ .json      │  │ 390 names  │  │
  │  └────────────────┘  └────────────────┘  └────────────┘  └────────────┘  │
  │  ┌────────────────┐  ┌────────────────┐  ┌────────────────────────────┐   │
  │  │  routes.json   │  │  TTL Cache     │  │  RideService (ride_service  │   │
  │  │  (train ref)   │  │  5min routes   │  │  .py) bike/tuk/car quotes  │   │
  │  │                │  │  60s disrupts  │  │                            │   │
  │  └────────────────┘  └────────────────┘  └────────────────────────────┘   │
  └──────────────────────────────────────────────────────────────────────────┘"""

    add_code_block(doc, arch_diagram, size=7.5)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Technology Stack", 1)

    tech_rows = [
        ("LangGraph",             ">=0.2.0",  "StateGraph orchestration, conditional edges, bounded loops"),
        ("LangChain",             ">=0.3.0",  "LLM abstraction, prompt management, retry handling"),
        ("langchain-groq",        ">=0.2.0",  "Groq LLM integration — llama-3.3-70b-versatile (free tier)"),
        ("googlemaps",            ">=4.10.0", "Google Maps Directions API — transit route discovery"),
        ("beautifulsoup4",        ">=4.12.0", "HTML parsing for trainschedule.lk scraping"),
        ("chromadb",              ">=0.5.0",  "Persisted vector store for route embeddings"),
        ("pydantic",              ">=2.7.0",  "Domain models, NLU models, API schemas — strict validation"),
        ("pydantic-settings",     ">=2.4.0",  "Settings management from .env + YAML"),
        ("streamlit",             ">=1.38.0", "Conversational chat UI"),
        ("fastapi",               ">=0.115.0","REST API layer"),
        ("uvicorn",               ">=0.30.0", "ASGI server for FastAPI"),
        ("gtts",                  ">=2.5.0",  "Google Text-to-Speech — per-message Read Aloud"),
        ("tenacity",              ">=9.0.0",  "LLM call retry logic (2 attempts on parse failure)"),
        ("httpx",                 ">=0.27.0", "Async HTTP client"),
        ("pyyaml",                ">=6.0.2",  "Config/prompts YAML loading"),
        ("rich",                  ">=13.8.0", "Terminal output formatting"),
        ("python-dotenv",         ">=1.0.0",  "Environment variable loading"),
        ("langchain-google-genai",">=2.0.0",  "Google Generative AI embeddings (text-embedding-004)"),
    ]
    make_table(doc, ["Package", "Version", "Purpose"], tech_rows, col_widths=[4, 2.5, 10.5])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. DOMAIN MODELS & ER DIAGRAM
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Domain Models & ER Diagram", 1)
    add_body(doc,
        "The following Entity-Relationship diagram shows all core domain entities, their "
        "attributes, and the relationships between them. Entities map directly to Pydantic "
        "models in the codebase.")

    er_diagram = """\
┌──────────────────────────────────────────────────────────────────────────────┐
│                         ENTITY-RELATIONSHIP DIAGRAM                           │
└──────────────────────────────────────────────────────────────────────────────┘

  ┌─────────────────────────┐          ┌──────────────────────────────────────┐
  │   ParsedIntent          │          │   RouteOption                         │
  │─────────────────────────│          │──────────────────────────────────────│
  │ PK  (session-scoped)    │          │ PK  route_id: str                    │
  │     language: Language  │◀────────▶│     line: str                        │
  │     origin: str|None    │ "informs"│     stops: list[str]                 │
  │     destination: str    │          │     departure_times: list[str]        │
  │     requested_time: str │          │     arrival_times: list[str]         │
  │     expected_arrival_   │          │     days_of_operation: list[str]     │
  │       time: str|None    │          │     transit_mode: TransitMode        │
  │     preferred_mode: str │          │     vehicle_type: str                │
  │     raw_intent: str     │          │     description: str                 │
  └────────────┬────────────┘          │     [origin property] = stops[0]     │
               │                       │     [destination property] = stops[-1]│
               │ "produces"            └──────────────────┬───────────────────┘
               │                                          │
               ▼                                          │ "is checked by"
  ┌──────────────────────────┐                            ▼
  │   AgentState             │          ┌──────────────────────────────────────┐
  │──────────────────────────│          │   DisruptionRecord                   │
  │     user_query           │          │──────────────────────────────────────│
  │     language             │◀────────▶│ PK  disruption_id: str              │
  │     origin               │ "holds"  │     train_id: str                   │
  │     destination          │          │     affected_segment: str            │
  │     requested_time       │          │     type: DisruptionType            │
  │     expected_arrival     │          │     delay_minutes: int|None         │
  │     preferred_mode       │          │     active: bool                    │
  │     candidate_routes     │          │     message: str                    │
  │     candidate_route      │          └──────────────────┬───────────────────┘
  │     ranked_routes        │                             │
  │     uber_options         │                             │ "produces"
  │     uber_last_mile       │                             ▼
  │     tts_audio            │          ┌──────────────────────────────────────┐
  │     disruption_status    │          │   DisruptionStatus                   │
  │     alternative_route    │          │──────────────────────────────────────│
  │     replan_attempts      │          │     level: DisruptionLevel          │
  │     final_response_*     │          │     disruption: DisruptionRecord     │
  │     trace                │          │       | None                        │
  │     error                │          │     [is_disrupted property]         │
  └──────────────────────────┘          └──────────────────────────────────────┘

  ┌──────────────────────────┐          ┌──────────────────────────────────────┐
  │   BusTimetable           │          │   Quote (RideService)                │
  │──────────────────────────│          │──────────────────────────────────────│
  │ PK  route_number: str    │          │ PK  quote_id: str                   │
  │     route_name: str      │          │     vehicle_type: str (bike|tuk|car)│
  │     origin_stop: str     │          │     label: str                      │
  │     destination_stop: str│          │     available: bool                 │
  │     departures: list[str]│          │     distance_km: float              │
  │     journey_time_min: int│          │     price: int (LKR)               │
  └──────────────────────────┘          │     currency: str                   │
                                        │     eta_min: int                    │
  ┌──────────────────────────┐          │     surge: float                    │
  │   Station                │          └──────────────────────────────────────┘
  │──────────────────────────│
  │ PK  name: str            │          ENUM REFERENCE
  │     (from stations.json  │          ─────────────
  │      390 official names  │          Language:      si | ta | en
  │      from               │          DisruptionType: delay | cancellation
  │      trainschedule.lk)  │          DisruptionLevel: clear | delayed | cancelled
  └──────────────────────────┘          TransitMode:   train | bus

  ┌─────────────────────────────────────────────────────────────────────────┐
  │  RELATIONSHIP SUMMARY                                                    │
  │  ParsedIntent  ────────▶ (informs) ─────▶  AgentState (many fields)     │
  │  AgentState    ────────▶ (holds) ──────▶  RouteOption[] (candidate_*)   │
  │  RouteOption   ────────▶ (checked by) ──▶  DisruptionRecord[]           │
  │  DisruptionRecord ──────▶ (produces) ───▶  DisruptionStatus             │
  │  BusTimetable  ────────▶ (enriches) ───▶  RouteOption (bus routes)     │
  │  Station       ────────▶ (maps to) ────▶  RouteOption (train routes)   │
  │  Quote         ────────▶ (supplements) ▶  AgentState (uber_options)    │
  └─────────────────────────────────────────────────────────────────────────┘"""

    add_code_block(doc, er_diagram, size=7)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5b — Model attribute tables
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5.1 RouteOption Model", 2)
    make_table(doc,
        ["Field", "Type", "Description"],
        [
            ("route_id",           "str",           "Unique identifier (e.g. GMAPS-0, TRAIN-sched0)"),
            ("line",               "str",           "Route line name (e.g. Colombo–Kandy)"),
            ("stops",              "list[str]",     "Ordered station/stop names"),
            ("departure_times",    "list[str]",     "HH:MM values, index-aligned with stops"),
            ("arrival_times",      "list[str]",     "HH:MM values, index-aligned with stops"),
            ("days_of_operation",  "list[str]",     "Service day names"),
            ("transit_mode",       "TransitMode",   "TRAIN | BUS"),
            ("vehicle_type",       "str",           "e.g. InterCity Express"),
            ("description",        "str",           "Human-readable multi-leg journey description"),
            ("origin [property]",  "str",           "stops[0] — computed property"),
            ("destination [prop]", "str",           "stops[-1] — computed property"),
        ], col_widths=[4, 3, 10])

    add_heading(doc, "5.2 DisruptionRecord Model", 2)
    make_table(doc,
        ["Field", "Type", "Description"],
        [
            ("disruption_id",    "str",          "Unique ID (D001, D002 ...)"),
            ("train_id",         "str",          "Affected train/service ID"),
            ("affected_segment", "str",          "Segment description for matching (e.g. 'Fort Railway station - Kandy')"),
            ("type",             "DisruptionType","delay | cancellation"),
            ("delay_minutes",    "int | None",   "Duration of delay (None for cancellation)"),
            ("active",           "bool",         "Whether this disruption is currently live"),
            ("message",          "str",          "User-facing explanation message"),
        ], col_widths=[4, 3, 10])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. AGENT GRAPH — NODE REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Agent Graph — Node Reference", 1)
    add_body(doc,
        "The LangGraph StateGraph contains 8 nodes executed in the order shown below. "
        "Each node reads a defined set of AgentState fields, performs its work, and writes "
        "results back into the state.")

    node_table = [
        ("1", "planner",    "Parse NLU intent; discover routes via Google Maps API",
         "user_query, origin, destination, requested_time, preferred_mode",
         "language, origin, destination, requested_time, expected_arrival_time, preferred_mode, candidate_routes, candidate_route",
         "Groq LLM (parse_query), Google Maps Directions API"),
        ("2", "bus_rag",    "Enrich bus routes with next scheduled departure from timetable",
         "candidate_routes, requested_time",
         "candidate_routes (updated), candidate_route (updated)",
         "data/bus_timetables.json (local file)"),
        ("3", "train_rag",  "Map station names; scrape trainschedule.lk for real train times",
         "candidate_routes, origin, destination, requested_time",
         "candidate_routes (enriched), candidate_route (updated)",
         "Groq LLM (map_station_name), trainschedule.lk (HTTP scrape)"),
        ("4", "ranker",     "Score & rank top-5 routes by deadline/time/mode/transfers",
         "candidate_routes, requested_time, expected_arrival_time",
         "ranked_routes, candidate_route (best ranked)",
         "None (pure logic)"),
        ("5", "uber",       "Get ride-hailing quotes when no transit fits or last-mile gap",
         "candidate_routes, ranked_routes, candidate_route, origin, destination, expected_arrival_time",
         "uber_options, uber_last_mile",
         "RideService (ride_service.py)"),
        ("6", "monitor",    "Check candidate/alternative route for active disruptions",
         "candidate_route, alternative_route, replan_attempts",
         "disruption_status",
         "data/disruptions.json (via TTL cache)"),
        ("7", "replanner",  "Select next untried route as alternative on disruption",
         "candidate_route, alternative_route, candidate_routes, replan_attempts",
         "alternative_route, replan_attempts (incremented)",
         "None (in-memory selection)"),
        ("8", "responder",  "Generate bilingual final response + TTS audio",
         "language, candidate_route, disruption_status, alternative_route, uber_options, uber_last_mile",
         "final_response_native, final_response_en, tts_audio",
         "Groq LLM (respond_*), gTTS (text-to-speech)"),
    ]

    for num, name, desc, reads, writes, external in node_table:
        add_heading(doc, f"Node {num}: {name.upper()}", 2)
        make_table(doc,
            ["Property", "Detail"],
            [
                ("Description",       desc),
                ("State Fields Read",  reads),
                ("State Fields Written", writes),
                ("External Services",  external),
            ], col_widths=[4.5, 12.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. USE CASE DIAGRAM
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Use Case Diagram", 1)
    add_body(doc,
        "The use case diagram identifies four actors and all major use cases the system supports. "
        "Relationships show includes (<<include>>) and extends (<<extend>>) dependencies.")

    uc_diagram = """\
┌────────────────────────────────────────────────────────────────────────────────────────┐
│                              COMMUTEAI — USE CASE DIAGRAM                               │
└────────────────────────────────────────────────────────────────────────────────────────┘

  ACTORS
  ──────
  [Commuter]         Sri Lanka resident planning a bus or train journey
  [Admin / Demo]     Demo operator who activates disruption scenarios
  [Groq LLM]         External AI service for NLU and response generation
  [Google APIs]      Google Maps Directions API + Google Text-to-Speech

                         ┌──────────────────────────────────────────┐
                         │           CommuteAI System               │
                         │                                          │
  ┌───────────┐          │  UC-01: Plan a Journey                   │
  │           │──────────│──────────────────────────────────────    │
  │           │          │     <<include>> UC-02                    │
  │           │          │     <<include>> UC-03                    │
  │           │          │     <<include>> UC-04                    │
  │           │          │                                          │
  │ COMMUTER  │          │  UC-02: Provide Location & Time          │
  │           │──────────│──────────────────────────────────────    │
  │           │          │     <<include>> UC-10 (multilingual)     │
  │           │          │                                          │
  │           │          │  UC-03: View Route Details               │
  │           │──────────│──────────────────────────────────────    │
  │           │          │     <<extend>> UC-05 (if disrupted)      │
  │           │          │     <<extend>> UC-06 (if no transit)     │
  │           │          │                                          │
  │           │          │  UC-04: Hear Response Aloud              │
  │           │──────────│──────────────────────────────────────    │
  │           │          │     <<include>> UC-11 (gTTS)             │
  │           │          │                                          │
  │           │          │  UC-05: Receive Disruption Alert         │
  │           │          │     <<include>> UC-07 (replanning)       │
  │           │          │                                          │
  │           │          │  UC-06: Get Ride-Hailing Quote           │
  │           │──────────│──────────────────────────────────────    │
  └───────────┘          │                                          │
                         │  UC-07: Auto-Replan with Alternative     │
  ┌───────────┐          │     <<extend>> UC-08 (if cap reached)    │
  │           │          │                                          │
  │  ADMIN /  │          │  UC-08: Receive No-Alternative Message   │
  │   DEMO    │          │                                          │
  │           │          │  UC-09: Activate Disruption Scenario     │
  │           │──────────│──────────────────────────────────────    │
  │           │          │                                          │
  └───────────┘          │  UC-10: Parse Multilingual Query         │
                         │     (Sinhala / Tamil / English)          │
  ┌───────────┐          │                                          │
  │           │          │  UC-11: Generate TTS Audio               │
  │ GROQ LLM  │──────────│──────────────────────────────────────    │
  │           │          │     UC-10, UC-02, UC-03 depend on this   │
  └───────────┘          │                                          │
                         │  UC-12: Rank Routes by Criteria          │
  ┌───────────┐          │     deadline / time / mode / transfers   │
  │  GOOGLE   │          │                                          │
  │  APIS     │──────────│  UC-13: Discover Transit Routes          │
  │           │          │     Google Maps Directions (transit)     │
  │           │──────────│  UC-14: Read Aloud Response              │
  └───────────┘          │     Google Text-to-Speech (gTTS)        │
                         └──────────────────────────────────────────┘"""

    add_code_block(doc, uc_diagram, size=7.5)

    add_heading(doc, "Use Case Descriptions", 2)
    uc_rows = [
        ("UC-01", "Plan a Journey",                 "Commuter",       "Commuter submits start/end location and time; system returns top route"),
        ("UC-02", "Provide Location & Time",        "Commuter",       "Conversational multi-turn intake: origin → destination → depart time → arrival deadline"),
        ("UC-03", "View Route Details",             "Commuter",       "See route card: mode, stops, departure/arrival, transfers, ranked alternatives"),
        ("UC-04", "Hear Response Aloud",            "Commuter",       "Click Read Aloud button → gTTS generates MP3 → audio auto-plays in browser"),
        ("UC-05", "Receive Disruption Alert",       "Commuter",       "System detects delay/cancellation and communicates it empathetically"),
        ("UC-06", "Get Ride-Hailing Quote",         "Commuter",       "When no transit fits, system offers bike/tuk/car quotes with price and ETA"),
        ("UC-07", "Auto-Replan with Alternative",   "Commuter",       "System selects next untried route and re-checks it (up to 2 attempts)"),
        ("UC-08", "Receive No-Alternative Message", "Commuter",       "When all options exhausted, system advises commuter to check later"),
        ("UC-09", "Activate Disruption Scenario",   "Admin / Demo",   "Demo operator activates D001/D002 via sidebar; cache invalidated"),
        ("UC-10", "Parse Multilingual Query",       "Groq LLM",       "parse_query/parse_update prompts detect language and extract intent fields"),
        ("UC-11", "Generate TTS Audio",             "Google APIs",    "gTTS converts English response text to MP3 bytes via Google TTS service"),
        ("UC-12", "Rank Routes by Criteria",        "System (Ranker)","Score routes: missed deadline, pre-requested, bus/train, transfers, arrival time"),
        ("UC-13", "Discover Transit Routes",        "Google APIs",    "Google Maps Directions API with mode=transit returns bus/train alternatives"),
        ("UC-14", "Scrape Train Schedules",         "trainschedule.lk","BeautifulSoup scrapes timetable HTML; Groq maps station names to official names"),
    ]
    make_table(doc, ["UC#", "Name", "Primary Actor", "Description"], uc_rows,
               col_widths=[1.5, 4, 3, 8.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. API REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. API Reference", 1)
    add_body(doc,
        "CommuteAI exposes a FastAPI REST API on port 8000 alongside the Streamlit UI. "
        "The API is primarily for integration with other systems or CI testing.")

    make_table(doc,
        ["Method", "Endpoint", "Request Body", "Response", "Description"],
        [
            ("POST", "/api/v1/query",                    "QueryRequest (user_query: str)",   "AgentResponse",   "Run full agent pipeline from raw query"),
            ("GET",  "/api/v1/health",                   "—",                               '{"status":"ok"}', "Health check"),
            ("POST", "/api/v1/disruptions/{id}/activate","—",                               '{"message":str}', "Invalidate cache (demo)"),
        ], col_widths=[1.5, 5, 4, 3, 4])

    add_heading(doc, "AgentResponse Schema", 2)
    make_table(doc,
        ["Field", "Type", "Description"],
        [
            ("user_query",             "str",          "Original query text"),
            ("language",               "str",          "Detected language: si | ta | en"),
            ("origin",                 "str | None",   "Parsed departure station"),
            ("destination",            "str | None",   "Parsed arrival station"),
            ("requested_time",         "str | None",   "HH:MM departure time"),
            ("candidate_route",        "dict | None",  "Best route (RouteOption serialized)"),
            ("disruption_status",      "dict | None",  "{'level': str, 'disruption': dict | None}"),
            ("alternative_route",      "dict | None",  "Alternative if disrupted"),
            ("final_response_native",  "str",          "Response in commuter's language"),
            ("final_response_en",      "str",          "English response (always present)"),
            ("trace",                  "list[str]",    "Ordered list of nodes that fired"),
            ("error",                  "str | None",   "Non-fatal error message if any"),
        ], col_widths=[4, 3, 10])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. DATA SCHEMA REFERENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Data Schema Reference", 1)

    add_heading(doc, "9.1 routes.json", 2)
    add_body(doc, "Train schedule reference data used for Chroma ingestion and keyword fallback retrieval.")
    add_code_block(doc, """\
[
  {
    "route_id":           "1015",
    "line":               "Colombo-Kandy",
    "stops":              ["Colombo Fort", "Maradana", "Kelaniya", ...],
    "departure_times":    ["06:30", "06:36", ...],
    "arrival_times":      ["06:30", "06:36", ...],
    "days_of_operation":  ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"],
    "transit_mode":       "train",
    "vehicle_type":       "InterCity Express",
    "description":        "Train 1015 — Colombo Fort to Kandy..."
  }
]""")

    add_heading(doc, "9.2 disruptions.json", 2)
    add_body(doc, "Simulated live disruption feed. Set active=true to trigger disruption flow in demos.")
    add_code_block(doc, """\
[
  {
    "disruption_id":    "D001",
    "train_id":         "GMAPS-0",
    "affected_segment": "Fort Railway station - Kandy",
    "type":             "delay",
    "delay_minutes":    45,
    "active":           true,
    "message":          "Train service between Fort and Kandy is delayed by 45 minutes..."
  }
]""")

    add_heading(doc, "9.3 bus_timetables.json", 2)
    add_body(doc, "Mock bus timetable data. Route numbers match Google Maps short_names (e.g. EX1, 346).")
    add_code_block(doc, """\
[
  {
    "route_number":         "EX1",
    "route_name":           "Colombo - Galle (Expressway SLTB)",
    "origin_stop":          "Bastian Mawatha Bus Stand",
    "destination_stop":     "Galle Highway And Long Distance",
    "departures":           ["05:00","05:30","06:00","06:30","07:00",...],
    "journey_time_minutes": 90
  }
]""")

    add_heading(doc, "9.4 stations.json", 2)
    add_body(doc, "390 official Sri Lanka Railways station names from trainschedule.lk. Used for station name mapping.")
    add_code_block(doc, """\
{
  "source": "https://trainschedule.lk/",
  "total": 390,
  "stations": ["Abanpola", "Ahangama", "Ahungalle", ..., "Yakkala"]
}""")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. CONFIGURATION & PROMPTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Configuration & Prompts", 1)

    add_heading(doc, "10.1 config/settings.yaml", 2)
    make_table(doc,
        ["Setting", "Default", "Description"],
        [
            ("retrieval.top_k",              "5",    "Maximum routes returned per query"),
            ("retrieval.similarity_threshold","0.60", "Semantic similarity cutoff (for ChromaDB when enabled)"),
            ("cache.enabled",                "true", "Enable/disable TTL response caching"),
            ("cache.ttl_seconds",            "300",  "Route cache TTL — 5 minutes"),
            ("graph.max_replan_attempts",    "2",    "Maximum replanning iterations before fallback to responder"),
        ], col_widths=[5.5, 2, 9.5])

    add_heading(doc, "10.2 config/prompts.yaml — Prompt Registry", 2)
    make_table(doc,
        ["Prompt Key", "Used By", "Purpose"],
        [
            ("parse_query",         "nlu/parser.py",      "Extract intent fields from raw multilingual query → JSON"),
            ("parse_update",        "nlu/parser.py",      "Merge follow-up message with existing intent context → JSON"),
            ("clarify_query",       "nlu/responder.py",   "Fallback clarification in 3 languages"),
            ("respond_clear",       "nlu/responder.py",   "Generate full bilingual response for undisrupted route"),
            ("respond_disrupted",   "nlu/responder.py",   "Generate bilingual response acknowledging disruption + alternative"),
            ("respond_no_alternative","nlu/responder.py", "Advise commuter when no alternatives remain"),
            ("map_station_name",    "graph/nodes/train_rag.py","Map Google Maps station name to official trainschedule.lk name"),
        ], col_widths=[4, 4, 9])

    add_heading(doc, "10.3 Environment Variables (.env)", 2)
    make_table(doc,
        ["Variable", "Required", "Description"],
        [
            ("GOOGLE_API_KEY",      "Yes", "Google Cloud API key (Gemini + Text-to-Speech + Maps fallback)"),
            ("GOOGLE_MAPS_API_KEY", "No",  "Dedicated Google Maps key (overrides GOOGLE_API_KEY for Maps calls)"),
            ("GROQ_API_KEY",        "Yes", "Groq API key — free tier, used for all LLM inference"),
            ("APP_ENV",             "No",  "deployment environment: development | production (default: development)"),
            ("LOG_LEVEL",           "No",  "Logging verbosity (default: INFO)"),
            ("MAX_REPLAN_ATTEMPTS", "No",  "Override max replan cap (default: 2)"),
            ("CACHE_TTL_SECONDS",   "No",  "Override route cache TTL in seconds (default: 300)"),
        ], col_widths=[4.5, 2, 10.5])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. NLU PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. NLU Pipeline", 1)
    add_body(doc,
        "The NLU pipeline converts raw multilingual commuter text into a structured ParsedIntent, "
        "then converts the agent result back into a bilingual human-readable response.")

    nlu_flow = """\
  RAW QUERY (any language: Sinhala / Tamil / English)
         │
         ▼
  sanitise_query()  ──  remove control chars, trim whitespace
         │
         ▼
  check_off_topic()  ──  keyword heuristics; raise OffTopicQueryError if not transport
         │
         ▼
  Groq LLM: parse_query prompt  ──  JSON extraction of 7 fields
         │
         ├─ language   (si | ta | en)
         ├─ origin     (English station name | null)
         ├─ destination (English station name | null)
         ├─ requested_time   (HH:MM | null)
         ├─ expected_arrival_time (HH:MM | null)
         ├─ preferred_mode (train | bus | null)
         └─ raw_intent  (brief English summary)
         │
         ▼
  ParsedIntent Pydantic model  ──  validators: normalise_time, normalise_station, normalise_mode
         │
         ▼
  Planner node reads ParsedIntent fields → populates AgentState
         │
         ▼
  [... graph execution ...]
         │
         ▼
  Groq LLM: respond_clear | respond_disrupted | respond_no_alternative
         │
         ├─ final_response_native  (response in commuter's language)
         └─ final_response_en      (English translation, always present)
         │
         ▼
  gTTS → tts_audio (MP3 bytes)  ──  Read Aloud button in UI"""

    add_code_block(doc, nlu_flow, size=8)

    add_heading(doc, "Multi-Turn Conversation", 2)
    add_body(doc,
        "In conversational mode, parse_update() merges each new message with accumulated intent. "
        "Rules: existing non-null fields are preserved unless the user explicitly changes them. "
        "Special phrases ('start over', 'no deadline', 'by train') are recognized in all 3 languages. "
        "The UI enforces an intake sequence: origin → destination → departure time → arrival deadline.")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. RIDE-HAILING INTEGRATION
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Ride-Hailing Integration", 1)
    add_body(doc,
        "When public transit cannot satisfy the commuter's request, CommuteAI falls back to "
        "ride-hailing quotes via a built-in RideService (ride_service.py at repo root). "
        "The quotes are deterministic (hash-seeded) so demos are reproducible.")

    make_table(doc,
        ["Vehicle", "Base Fare (LKR)", "Per-km Rate (LKR)", "Base ETA (min)"],
        [
            ("Bike (motorbike/scooter)", "60",  "45",  "3"),
            ("Tuk-tuk (trishaw/auto)",   "100", "75",  "5"),
            ("Car (taxi/cab/sedan)",      "220", "120", "7"),
        ], col_widths=[5, 3.5, 3.5, 3.5])

    add_heading(doc, "Trigger Conditions", 2)
    make_table(doc,
        ["Condition", "Action", "State Field"],
        [
            ("No transit routes found (candidate_routes empty)",       "Full-journey Uber quotes: origin → destination",    "uber_options"),
            ("All ranked routes miss arrival deadline",                 "Full-journey Uber quotes: origin → destination",    "uber_options"),
            ("Transit found but route has last-mile walking gap",      "Last-mile quotes: final stop → destination",        "uber_last_mile"),
            ("Good transit found, no gap",                             "No Uber — transit alone is sufficient",             "(none)"),
        ], col_widths=[6, 6, 3])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 13. UI & USER EXPERIENCE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13. UI & User Experience", 1)
    add_body(doc,
        "The Streamlit chat UI provides a fully conversational experience with a sidebar for "
        "context display and demo controls. The main area hosts the chat interface with structured "
        "result components rendered below each assistant message.")

    add_heading(doc, "Conversational Intake Flow", 2)
    intake_flow = """\
  Turn 1: User says "I want to go to Kandy"
    → parse_query detects destination=Kandy, origin=None
    → Bot asks: "Departing from [origin]? Where will you be departing from?"

  Turn 2: User says "From Colombo"
    → parse_update detects origin=Colombo
    → Both fields collected; Bot asks: "What time are you planning to depart?"

  Turn 3: User says "8am" or "now"
    → requested_time = "08:00" or None
    → Bot asks: "Do you have an arrival deadline?"

  Turn 4: User says "No" or "before 10am"
    → expected_arrival_time = None or "10:00"
    → All fields collected → run_commute_agent_from_intent() called
    → Full route plan displayed"""
    add_code_block(doc, intake_flow, size=8)

    add_heading(doc, "Result Components", 2)
    make_table(doc,
        ["Component", "Description", "When Shown"],
        [
            ("Route Card",            "Mode, line, stops, departure/arrival times",                     "Always when route found"),
            ("Disruption Banner",     "Red/orange alert showing disruption type and delay",              "When disruption detected"),
            ("Alternative Route Card","Second route card for the replanned alternative",                "When disruption + alternative found"),
            ("Ranked Routes Expander","Collapsible list of top-5 scored routes",                        "Always when >1 route exists"),
            ("Uber Card",             "3-column metric: bike / tuk / car with price and ETA",          "When Uber needed (no transit / last-mile)"),
            ("Read Aloud Button",     "Generates and plays gTTS MP3 for English response text",        "Every assistant message"),
            ("English Translation",   "Collapsible expander showing English version of native response","When language != en"),
            ("Agent Trace",           "Ordered list of nodes that executed with icons and description","Every response"),
            ("Agent Note Expander",   "Non-fatal warning/error from agent",                            "When state.error is set"),
        ], col_widths=[4, 7, 5])

    add_heading(doc, "Sidebar Features", 2)
    add_bullet(doc, "Current Journey panel — shows accumulated origin, destination, time, mode, deadline")
    add_bullet(doc, "New Journey button — resets all session state")
    add_bullet(doc, "Demo Controls — Activate delay (D001), Activate cancellation (D002), Clear disruptions")
    add_bullet(doc, "Example queries in English, Sinhala, and Tamil")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 14. EXECUTION FLOW WALKTHROUGH
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14. Execution Flow Walkthrough", 1)

    add_heading(doc, "14.1 Happy Path — Clear Train Route", 2)
    happy = """\
  Input: "I want to take a train from Colombo to Kandy at 8am"
  ──────────────────────────────────────────────────────────────────────
  [PLANNER]
    Groq NLU: language=en, origin=Colombo, destination=Kandy, time=08:00, mode=train
    Google Maps: 3 transit routes found (2 train, 1 bus)
    Filter by mode=train → 2 train routes
    candidate_routes=[route_A, route_B], candidate_route=route_A
    trace: ["planner"]

  [BUS_RAG]
    No bus routes in candidate_routes → nothing to enrich
    trace: ["planner", "bus_rag"]

  [TRAIN_RAG]
    Groq LLM maps "Colombo" → "Colombo Fort"
    Groq LLM maps "Kandy" → "Kandy"
    HTTP scrape: trainschedule.lk/schedule/abucnia/colombo-fort-to-kandy-train-timetable
    5 trains found departing after 08:00: 08:15, 09:30, 10:00, 11:15, 13:00
    Replace route_A, route_B with 5 enriched entries
    trace: ["planner", "bus_rag", "train_rag"]

  [RANKER]
    Score all 5 routes: none miss deadline, all depart after 08:00, all trains
    Rank by earliest arrival → 08:15 departure wins
    ranked_routes=[route_0815, ...], candidate_route=route_0815
    trace: ["planner", "bus_rag", "train_rag", "ranker"]

  [UBER]
    Good transit found, no last-mile gap
    uber_options=None, uber_last_mile=None
    trace: ["planner", "bus_rag", "train_rag", "ranker", "uber"]

  [MONITOR]
    Check route_0815: origin="Colombo Fort", destination="Kandy"
    Load disruptions.json → D001 affected_segment="Fort Railway station - Kandy"
    Wait: D001.active=false in this run → no match
    disruption_status={level:"clear", disruption:None}
    trace: ["planner", "bus_rag", "train_rag", "ranker", "uber", "monitor"]

  → _route_after_monitor: level==CLEAR → RESPONDER

  [RESPONDER]
    generate_response: respond_clear template
    Groq LLM: "Great news! The 08:15 Colombo Fort to Kandy train..."
    gTTS: MP3 audio bytes
    final_response_native=<English>, final_response_en=<English>, tts_audio=<bytes>
    trace: ["planner", "bus_rag", "train_rag", "ranker", "uber", "monitor", "responder"]

  Output: Journey plan + Read Aloud button + Ranked routes expander + Agent trace"""
    add_code_block(doc, happy, size=8)

    add_heading(doc, "14.2 Disruption Path — Delay Detected, Alternative Found", 2)
    disrupted = """\
  D001 set to active=true in disruptions.json
  ──────────────────────────────────────────────────────────────────────
  [MONITOR — Iteration 1]
    Check candidate_route (Colombo Fort → Kandy route_0815)
    D001 active: affected_segment contains "Fort Railway station - Kandy"
    Match found: type=delay → level=DELAYED
    disruption_status={level:"delayed", disruption:D001}

  → _route_after_monitor: level!=CLEAR, replan_attempts=0 → REPLANNER

  [REPLANNER]
    Tried routes: {route_0815.route_id}
    Remaining: [route_0930, route_1000, ...]
    alternative_route = route_0930
    replan_attempts = 1

  → _route_after_replanner: always → MONITOR

  [MONITOR — Iteration 2]
    replan_attempts=1 → check alternative_route (route_0930)
    D001 affected_segment: "Fort Railway station - Kandy" still matches
    disruption_status={level:"delayed", disruption:D001}

  → replan_attempts=1 >= max(2)? No → REPLANNER again

  [REPLANNER — 2nd attempt]
    Tried: {route_0815, route_0930}
    Remaining: [route_1000, ...]
    alternative_route = route_1000
    replan_attempts = 2

  [MONITOR — Iteration 3]
    Check route_1000 → D001 match → still delayed
    replan_attempts=2 >= max(2) → RESPONDER (cap reached)

  [RESPONDER]
    disruption present, alternative_route found
    generate_response: respond_disrupted template
    "We're sorry — the Colombo Fort–Kandy service is delayed 45 min.
     We recommend the 10:00 service as your best alternative..."
    tts_audio = MP3 of English response

  Output: Disruption banner + original + alternative route cards + Read Aloud"""
    add_code_block(doc, disrupted, size=8)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 15. DESIGN DECISIONS & TRADE-OFFS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "15. Design Decisions & Trade-offs", 1)

    decisions = [
        ("Google Maps for Route Discovery",
         "Using the Google Maps Directions API (transit mode) means we get real Sri Lanka bus route "
         "numbers, stop names, and multi-leg transfers without maintaining our own timetable database. "
         "Trade-off: API quota limits and requires an active internet connection.",
         "The API returns line.short_name with actual Sri Lanka route numbers (346-1, EX1-18) making "
         "the bus_rag matching straightforward."),
        ("Groq LLM (Free Tier) for NLU",
         "Groq's free tier with llama-3.3-70b-versatile provides fast, cost-free inference for NLU "
         "and response generation. Trade-off: rate limits, and free tier may have lower availability.",
         "All LLM calls wrapped in Tenacity retry (2 attempts) for resilience."),
        ("Bus Disruption: CLEAR Bypass",
         "Google Maps does not expose bus disruption data via its API. We chose to always mark bus "
         "routes as CLEAR and skip disruption checking rather than fabricate data.",
         "Documented clearly in monitor node; bus routes always pass disruption check."),
        ("Segment-Based Disruption Matching",
         "Disruptions are matched by checking whether route origin/destination appears in the "
         "affected_segment string rather than exact train ID matching. This works for both "
         "manual timetable routes and Google Maps routes.",
         "Trade-off: may produce false positives if station name appears in unrelated segment."),
        ("trainschedule.lk Scraping",
         "Real train schedules are scraped on-demand rather than stored in a database. "
         "Trade-off: network dependency and potential breakage if the website structure changes.",
         "Graceful fallback: if scraping fails, train routes retain Google Maps estimated times."),
        ("Bounded Replan Loop",
         "The replanning loop is capped at MAX_REPLAN_ATTEMPTS (default 2) enforced at both the "
         "graph routing function and the replanner node. Double guard prevents infinite loops.",
         "Explained as a deliberate defense-in-depth design to guarantee graph termination."),
        ("Conversational vs. Single-Turn",
         "The UI accumulates intent over multiple turns rather than requiring a single complete "
         "query. This makes the agent more accessible to users who may not know how to structure "
         "a complete transit query in one message.",
         "Trade-off: more complex state management in the UI layer."),
        ("ChromaDB Not Yet Wired",
         "The ChromaDB vector store and embedding pipeline exist in the codebase but the embedder "
         "requires a real GOOGLE_API_KEY. The bus_rag and retrieval.py both fall back to JSON "
         "keyword matching, which is sufficient for the demo.",
         "The full semantic search path is ready to activate once a production API key is set."),
    ]

    for title, explanation, tradeoff in decisions:
        add_heading(doc, title, 3)
        add_body(doc, explanation)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"Trade-off / Note: {tradeoff}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_ACCENT

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 16. DEPLOYMENT GUIDE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "16. Deployment Guide", 1)

    add_heading(doc, "Prerequisites", 2)
    for req in ["Python 3.11+", "uv package manager (pip install uv)", "GROQ_API_KEY (free at console.groq.com)", "GOOGLE_API_KEY (optional — for Maps and TTS)"]:
        add_bullet(doc, req)

    add_heading(doc, "Installation & Setup", 2)
    add_code_block(doc, """\
# 1. Clone the repository
git clone <repo-url>
cd AGENTRIX26-TEAM23-AlphaZero

# 2. Install dependencies
uv sync

# 3. Configure environment
cp .env.example .env
# Edit .env: set GROQ_API_KEY and GOOGLE_API_KEY

# 4. (Optional) Ingest routes into ChromaDB
uv run ingest""")

    add_heading(doc, "Running the Application", 2)
    add_code_block(doc, """\
# Streamlit UI (primary interface)
uv run streamlit run src/commute_agent/ui/app.py
# Opens at: http://localhost:8501

# FastAPI backend (optional)
uv run uvicorn commute_agent.api.main:app --reload --host 0.0.0.0 --port 8000
# API docs at: http://localhost:8000/docs

# CLI smoke test
uv run run-agent""")

    add_heading(doc, "Demo Disruption Controls", 2)
    add_body(doc,
        "In the Streamlit sidebar, use the Demo Controls panel to simulate disruptions "
        "without manual file editing. Click 'Activate: Train Delay (45 min)' to set D001 active, "
        "then run a Colombo–Kandy query to see the full disruption + replanning flow. "
        "Click 'Clear all disruptions' to reset.")

    add_heading(doc, "Project Scripts", 2)
    make_table(doc,
        ["Command", "Equivalent", "Description"],
        [
            ("uv run ingest",         "python -m commute_agent.rag.ingest",           "Ingest routes.json into ChromaDB"),
            ("uv run run-agent",      "python -m commute_agent.graph.builder",        "CLI smoke test"),
            ("uv run streamlit ...",  "streamlit run src/commute_agent/ui/app.py",    "Launch UI"),
            ("uv run uvicorn ...",    "uvicorn commute_agent.api.main:app",           "Launch API"),
        ], col_widths=[4, 5.5, 7.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # APPENDIX — AgentState Field Reference
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Appendix A: Complete AgentState Field Reference", 1)
    make_table(doc,
        ["Field", "Type", "Written By", "Read By", "Description"],
        [
            ("user_query",             "str",           "UI/caller",  "planner",        "Raw commuter query text"),
            ("language",               "str",           "planner",    "responder, UI",  "Detected language: si|ta|en"),
            ("origin",                 "str|None",      "planner",    "train_rag, uber","Departure station (English)"),
            ("destination",            "str|None",      "planner",    "train_rag, uber","Arrival station (English)"),
            ("requested_time",         "str|None",      "planner",    "bus_rag, train_rag, ranker","Departure HH:MM"),
            ("expected_arrival_time",  "str|None",      "planner",    "ranker, uber",   "Arrival deadline HH:MM"),
            ("preferred_mode",         "str|None",      "planner",    "planner (filter)","train|bus|None"),
            ("candidate_routes",       "list[dict]",    "planner",    "bus_rag, train_rag, ranker, uber, replanner","All routes from Google Maps"),
            ("candidate_route",        "dict|None",     "planner / ranker","monitor, replanner, responder, UI","Best single route"),
            ("ranked_routes",          "list[dict]",    "ranker",     "uber, UI",       "Top-5 scored routes"),
            ("uber_options",           "list[dict]|None","uber",      "responder, UI",  "Full-trip ride quotes"),
            ("uber_last_mile",         "list[dict]|None","uber",      "responder, UI",  "Last-mile ride quotes"),
            ("tts_audio",              "bytes|None",    "responder",  "UI (state only)","MP3 audio (backend; UI uses on-demand gTTS)"),
            ("disruption_status",      "dict|None",     "monitor",    "builder routing, responder","Disruption level + record"),
            ("alternative_route",      "dict|None",     "replanner",  "monitor (itr2+), responder, UI","Alternative route"),
            ("replan_attempts",        "int",           "replanner",  "builder routing, monitor","Iteration counter"),
            ("final_response_native",  "str",           "responder",  "UI",             "Response in commuter's language"),
            ("final_response_en",      "str",           "responder",  "UI",             "English response (always present)"),
            ("trace",                  "list[str]",     "every node", "UI",             "Ordered node execution log"),
            ("error",                  "str|None",      "planner / responder","responder, UI","Non-fatal error message"),
        ], col_widths=[4, 2.5, 2.5, 3, 5])

    doc.save(OUTPUT)
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    build_report()
